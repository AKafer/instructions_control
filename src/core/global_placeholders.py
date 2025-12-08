from copy import deepcopy
from typing import List, Dict

from sqlalchemy import select

from database.models import Config
from database.orm import Session
from docx import Document

PROFESSION = '{{профессия}}'
POINT_NUMBER = '{{номер пп}}'
INSTRUCTION_NAME = '{{наименование инструкции}}'
INSTRUCTION_USER_NAME = '{{наименование инструкции работника}}'
INSTRUCTION_NUMBER = '{{номер инструкции}}'
NPA_SIZ = '{{пункт 767}}'
QUANTITY_SIZ = '{{кол СИЗ}}'
NAME_SIZ = '{{Наименование СИЗ}}'
START_DATE_SIZ = '{{дата выдачи СИЗ}}'
UNIT_OF_MEASUREMENT_SIZ = '{{шт-пар}}'
NON_QUALIFY_PROF = '{{профессии освобожденные от первичного инструктажа}}'


class DocumentCreateError(Exception):
    pass


def doc_contains_placeholder(doc, placeholder) -> bool:
    for p in doc.paragraphs:
        if placeholder in p.text:
            return True

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if placeholder in p.text:
                        return True

    for section in doc.sections:
        for header_footer in [section.header, section.footer]:
            for p in header_footer.paragraphs:
                if placeholder in p.text:
                    return True
            for table in header_footer.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            if placeholder in p.text:
                                return True

    return False


async def get_placeholders() -> dict[str, str]:
    async with Session() as session:
        query = select(Config).where(Config.id == 1)
        config = await session.scalar(query)
        placeholders = config.placeholders
        global_placeholders = placeholders.get('global_placeholders', {})
        placeholders_dict = {}
        for key, item in global_placeholders.items():
            placeholders_dict[item['key']] = item['value']
        return placeholders_dict


def replace_in_paragraph_runs(
    paragraph, placeholder: str, replacement: str
) -> bool:
    if not placeholder:
        return False

    runs = paragraph.runs
    if not runs:
        return False

    full_text = ''.join(r.text for r in runs)
    idx = full_text.find(placeholder)
    if idx == -1:
        return False

    start_pos = idx
    end_pos = idx + len(placeholder)

    accum = 0
    start_run_idx = None
    end_run_idx = None
    start_offset = None
    end_offset = None

    for i, r in enumerate(runs):
        run_len = len(r.text)
        if start_run_idx is None and accum + run_len > start_pos:
            start_run_idx = i
            start_offset = start_pos - accum
        if end_run_idx is None and accum + run_len >= end_pos:
            end_run_idx = i
            end_offset = end_pos - accum
            break
        accum += run_len

    if start_run_idx is None or end_run_idx is None:
        return False

    if start_run_idx == end_run_idx:
        run = runs[start_run_idx]
        run.text = (
            run.text[:start_offset] + replacement + run.text[end_offset:]
        )
        return True

    first_run = runs[start_run_idx]
    before = first_run.text[:start_offset]
    last_run = runs[end_run_idx]
    after = last_run.text[end_offset:]

    first_run.text = before + replacement + after

    for j in range(start_run_idx + 1, end_run_idx + 1):
        runs[j].text = ''

    return True


def replace_in_paragraphs(
    doc: Document, placeholders: Dict[str, str]
) -> Document:
    for p in doc.paragraphs:
        for ph, val in placeholders.items():
            replace_in_paragraph_runs(p, ph, val)
    return doc


def replace_placeholders_in_cell(cell, placeholders: Dict[str, str]):
    for p in cell.paragraphs:
        for ph, val in placeholders.items():
            replace_in_paragraph_runs(p, ph, val)


def replace_in_tables(doc: Document, placeholders: Dict[str, str]) -> Document:
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                replace_placeholders_in_cell(cell, placeholders)
    return doc


def replace_in_headers_footers(
    doc: Document, placeholders: Dict[str, str]
) -> Document:
    for section in doc.sections:
        for header_footer in (section.header, section.footer):
            for p in header_footer.paragraphs:
                for ph, val in placeholders.items():
                    replace_in_paragraph_runs(p, ph, val)
            for table in header_footer.tables:
                for row in table.rows:
                    for cell in row.cells:
                        replace_placeholders_in_cell(cell, placeholders)
    return doc


async def replace_global_placeholders_in_doc(doc):
    placeholders = await get_placeholders()
    doc = replace_in_paragraphs(doc, placeholders)
    doc = replace_in_tables(doc, placeholders)
    doc = replace_in_headers_footers(doc, placeholders)
    return doc


async def fill_template_placeholders(doc, placeholders):
    """placeholders: list of Placeholder models
    {'key': str, 'value': str}
    """
    placeholders_dict = {ph.key: ph.value for ph in placeholders}
    doc = replace_in_paragraphs(doc, placeholders_dict)
    doc = replace_in_tables(doc, placeholders_dict)
    doc = replace_in_headers_footers(doc, placeholders_dict)
    return doc


def fill_table_with_items(
    doc: Document,
    items: List[Dict[str, str]],
    required_placeholders: List[str],
) -> Document:
    target_table = None
    for t in doc.tables:
        for row in t.rows:
            row_text = '\n'.join(c.text for c in row.cells)
            if all(ph in row_text for ph in required_placeholders):
                target_table = t
                break
        if target_table:
            break

    if not target_table:
        raise DocumentCreateError(
            f'Not found target table containing placeholders {required_placeholders}'
        )

    template_row = None
    template_index = None
    for i, row in enumerate(target_table.rows):
        text = '\n'.join(cell.text for cell in row.cells)
        if all(ph in text for ph in required_placeholders):
            template_row = row
            template_index = i
            break

    if not template_row:
        raise DocumentCreateError('Template row not found in the table.')

    while len(target_table.rows) > template_index + 1:
        target_table._tbl.remove(target_table.rows[-1]._tr)

    for counter, item_values in enumerate(items, start=1):
        new_tr = deepcopy(template_row._tr)
        target_table._tbl.append(new_tr)
        row = target_table.rows[-1]

        for cell in row.cells:
            for p in cell.paragraphs:
                replace_in_paragraph_runs(p, POINT_NUMBER, str(counter))
                for ph, val in item_values.items():
                    replace_in_paragraph_runs(p, ph, str(val))

    target_table._tbl.remove(template_row._tr)
    return doc
