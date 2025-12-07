import logging
from datetime import datetime, date
from io import BytesIO
from typing import Any, Sequence, List
from zipfile import ZipFile, ZIP_DEFLATED

from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession
from docx import Document

from constants import personal_placeholders, FileTemplatesNamingEnum
from core.global_placeholders import (
    replace_global_placeholders_in_doc,
    fill_template_placeholders,
    replace_in_paragraphs,
    replace_in_tables,
    replace_in_headers_footers,
    DocumentCreateError,
    POINT_NUMBER,
    INSTRUCTION_USER_NAME,
    NAME_SIZ,
    START_DATE_SIZ,
    NPA_SIZ,
    QUANTITY_SIZ,
    UNIT_OF_MEASUREMENT_SIZ,
    INSTRUCTION_NAME,
    INSTRUCTION_NUMBER,
    fill_table_with_items,
    PROFESSION,
    NON_QUALIFY_PROF,
)
from database.models import User, Documents, DocumentTypes, Instructions
from database.orm import Session
from web.documents.schemas import CreateDocument, Placeholder


logger = logging.getLogger('control')


async def update_document_db(
    document: Documents, **update_data: dict
) -> Documents:
    for field, value in update_data.items():
        setattr(document, field, value)
    return document


async def check_document_create(
    db_session: AsyncSession, document_input: CreateDocument
) -> None:
    query = select(User).filter(User.id == document_input.user_id)
    user = await db_session.scalar(query)
    if user is None:
        raise DocumentCreateError('User not found')
    query = select(DocumentTypes).filter(
        DocumentTypes.id == document_input.document_type_id
    )
    document_type = await db_session.scalar(query)
    if document_type is None:
        raise DocumentCreateError('Document type not found')


def replace_list_placeholders_in_doc(
    doc: Document,
    items: List[str],
) -> Document:
    new_items = []
    for item in items:
        new_items.append({NON_QUALIFY_PROF: item})
    placeholders = [NON_QUALIFY_PROF]
    doc = fill_table_with_items(
        doc, items=new_items, required_placeholders=placeholders
    )

    return doc


def replace_ins_generate_in_doc(doc, sections, profession) -> Document:
    name_mapping = {
        'general': 'Часть 1',
        'before': 'Часть 2',
        'during': 'Часть 3',
        'accidents': 'Часть 4',
        'after': 'Часть 5',
    }
    for paragraph in doc.paragraphs:
        if PROFESSION in paragraph.text:
            paragraph.text = paragraph.text.replace(PROFESSION, profession)

        for key, part_name in name_mapping.items():
            if '{{' + part_name + '}}' in paragraph.text:
                section_text = sections.get(key)
                if section_text:
                    paragraph.text = paragraph.text.replace(
                        '{{' + part_name + '}}', section_text.get('text', '')
                    )
                else:
                    paragraph.text = paragraph.text.replace(
                        '{{' + part_name + '}}', ''
                    )

    return doc


async def generate_document_in_memory(
    template_path: str, callback: replace_ins_generate_in_doc, **kwargs
) -> BytesIO:
    doc = Document(template_path)
    callback_map = {
        'replace_ins_generate_in_doc': replace_ins_generate_in_doc,
        'replace_list_placeholders_in_doc': replace_list_placeholders_in_doc,
    }
    local_placeholders_replace = callback_map.get(callback)
    doc = local_placeholders_replace(doc, **kwargs)
    doc = await replace_global_placeholders_in_doc(doc)
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def get_nested_value(obj: Any, path: str) -> Any:
    parts = path.split('.')
    current = obj
    try:
        for part in parts:
            if current is None:
                return ''
            if isinstance(current, dict):
                current = current.get(part, None)
            else:
                current = getattr(current, part, None)
        if isinstance(current, (datetime, date)):
            return current.strftime('%d.%m.%Y')
        return current if current is not None else ''
    except Exception as e:
        logger.error(f'Error getting nested value for path "{path}": {e}')
        return ''


async def fill_user_data_in_doc(doc: Document, user: Any) -> Document:
    placeholders_dict = {}
    for item in personal_placeholders:
        key = item['key']
        path = item['value']
        value = get_nested_value(user, path)
        placeholders_dict[key] = str(value) if value is not None else '-'

    replace_in_paragraphs(doc, placeholders_dict)
    replace_in_tables(doc, placeholders_dict)
    replace_in_headers_footers(doc, placeholders_dict)

    return doc


async def fill_user_tables_ins_in_doc(doc: Document, user: User) -> Document:
    ins_titles = [
        {INSTRUCTION_USER_NAME: ins.title} for ins in user.instructions
    ]
    doc = fill_table_with_items(
        doc,
        items=ins_titles,
        required_placeholders=[POINT_NUMBER, INSTRUCTION_USER_NAME],
    )

    return doc


async def insert_instruction_list_in_doc(doc: Document) -> Document:
    async with Session() as session:
        result = await session.execute(select(Instructions))
        instructions = result.scalars().all()

    items = [
        {INSTRUCTION_NAME: instr.title, INSTRUCTION_NUMBER: instr.number}
        for instr in instructions
    ]

    return fill_table_with_items(
        doc,
        items=items,
        required_placeholders=[POINT_NUMBER, INSTRUCTION_NAME],
    )


async def fill_user_tables_norm_siz_in_doc(doc: Document, user) -> Document:
    material_norm_types = (
        user.profession.norm.material_norm_types or []
        if user.profession.norm
        else []
    )

    items = []
    for counter, mat in enumerate(material_norm_types, start=1):
        material_type = mat.material_type
        items.append(
            {
                NAME_SIZ: str(material_type.title),
                NPA_SIZ: str(mat.npa_link),
                QUANTITY_SIZ: str(
                    int(mat.quantity) if mat.quantity is not None else '-'
                ),
                UNIT_OF_MEASUREMENT_SIZ: getattr(
                    material_type.unit_of_measurement, 'value', ''
                )
                if mat.quantity
                else '',
            }
        )

    if not items:
        items.append(
            {
                NAME_SIZ: '-',
                NPA_SIZ: '-',
                QUANTITY_SIZ: '-',
                UNIT_OF_MEASUREMENT_SIZ: '',
            }
        )

    return fill_table_with_items(
        doc, items=items, required_placeholders=[POINT_NUMBER, NAME_SIZ]
    )


async def fill_user_tables_fact_siz_in_doc(doc: Document, user) -> Document:
    material_norm_types = (
        user.profession.norm.material_norm_types or []
        if user.profession.norm
        else []
    )

    items = []
    for counter, mat in enumerate(material_norm_types, start=1):
        material_type = getattr(mat, 'material_type', None)
        items.append(
            {
                NAME_SIZ: material_type.title,
                QUANTITY_SIZ: str(
                    int(mat.quantity) if mat.quantity is not None else '-'
                ),
                UNIT_OF_MEASUREMENT_SIZ: getattr(
                    material_type.unit_of_measurement, 'value'
                )
                if mat.quantity
                else '',
                START_DATE_SIZ: START_DATE_SIZ,
            }
        )

    if not items:
        items.append(
            {
                NAME_SIZ: '-',
                QUANTITY_SIZ: '-',
                UNIT_OF_MEASUREMENT_SIZ: '-',
                START_DATE_SIZ: '-',
            }
        )

    return fill_table_with_items(
        doc, items=items, required_placeholders=[POINT_NUMBER, NAME_SIZ]
    )


async def generate_zip_personals_in_memory(
    template: str,
    template_path: str,
    users_data: Sequence[User],
    placeholders: list[Placeholder],
) -> BytesIO:
    zip_buffer = BytesIO()

    with ZipFile(zip_buffer, mode='w', compression=ZIP_DEFLATED) as zip_file:
        for user in users_data:
            doc = Document(template_path)

            if template == FileTemplatesNamingEnum.LNNA_ACK:
                doc = await fill_user_tables_ins_in_doc(doc, user)

            if template == FileTemplatesNamingEnum.LK_SIZ:
                doc = await fill_user_tables_norm_siz_in_doc(doc, user)
                doc = await fill_user_tables_fact_siz_in_doc(doc, user)

            doc = await fill_user_data_in_doc(doc, user)
            doc = await fill_template_placeholders(doc, placeholders)
            doc = await replace_global_placeholders_in_doc(doc)

            doc_buffer = BytesIO()
            doc.save(doc_buffer)
            doc_buffer.seek(0)

            user_file_name = f'{user.last_name} {user.name}'
            filename_in_zip = f'{user_file_name }.docx'
            zip_file.writestr(filename_in_zip, doc_buffer.read())

    zip_buffer.seek(0)
    return zip_buffer


async def generate_doc_organization_in_memory(
    template: str,
    template_path: str,
    placeholders: list[Placeholder],
) -> BytesIO:
    doc = Document(template_path)

    if template in [
        FileTemplatesNamingEnum.LIST_INSTRUCTIONS,
        FileTemplatesNamingEnum.ORDER_APPROVE_LNA,
    ]:
        doc = await insert_instruction_list_in_doc(doc)

    doc = await fill_template_placeholders(doc, placeholders)
    doc = await replace_global_placeholders_in_doc(doc)

    output_buffer = BytesIO()
    doc.save(output_buffer)
    output_buffer.seek(0)
    return output_buffer

