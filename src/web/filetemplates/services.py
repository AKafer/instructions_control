import logging
import os
import re
import unicodedata

import aiofiles as aiof
from docx import Document
from fastapi import Request, UploadFile
from settings import BASE_URL, STATIC_FOLDER, TEMPLATES_DIR, TEMPLATES_FOLDER
from templates_config import FileTemplatesNamingEnum
from web.filetemplates.schemas import DocumentField, RequestModel

logger = logging.getLogger('control')


def is_allowed_name(filename: str) -> bool:
    try:
        FileTemplatesNamingEnum(filename)
        return True
    except ValueError:
        return False


async def save_file(new_file: UploadFile, title) -> str:
    _, suffix = os.path.splitext(new_file.filename)
    filename = f'{title}{suffix}'
    path_to_file = os.path.join(TEMPLATES_DIR, filename)
    async with aiof.open(path_to_file, 'wb+') as f:
        await f.write(new_file.file.read())
    return filename


def get_full_link(request: Request, filename: str) -> str:
    for file in os.listdir(TEMPLATES_DIR):
        if file.startswith(filename):
            filename = file
            break
    base_url = BASE_URL or str(request.base_url)
    return f'{base_url}api/{STATIC_FOLDER}/{TEMPLATES_FOLDER}/{filename}'


def rename_old_file(filename: str) -> (str | None, int):
    old_name = None
    renamed_count = 0
    for file in os.listdir(TEMPLATES_DIR):
        if file.startswith(filename):
            renamed_count += 1
            old_name = f'old--{file}'
            os.rename(
                os.path.join(TEMPLATES_DIR, file),
                os.path.join(TEMPLATES_DIR, old_name),
            )
    return old_name, renamed_count


def roolback_rename_file(filename: str) -> None:
    try:
        for file in os.listdir(TEMPLATES_DIR):
            if file.startswith(f'old--{filename}'):
                new_name = file.split('--')[-1]
                os.rename(
                    os.path.join(TEMPLATES_DIR, file),
                    os.path.join(TEMPLATES_DIR, new_name),
                )
    except (FileNotFoundError, OSError):
        pass


def delete_file(filename: str) -> None:
    logger.debug(f'Delete file {filename}')
    try:
        for file in os.listdir(TEMPLATES_DIR):
            if file.startswith(filename):
                os.remove(os.path.join(TEMPLATES_DIR, file))
    except FileNotFoundError:
        pass


def replace_text_in_runs(runs, replacements):
    full_text = ''.join(run.text for run in runs).strip()
    full_text = unicodedata.normalize('NFKD', full_text)
    full_text = re.sub(r'\s+', ' ', full_text)
    full_text = full_text.replace('“', '{').replace('”', '}')

    if not full_text:
        return

    modified = False
    for key, value in replacements.items():
        search_key = f'{{{{{key}}}}}'
        if search_key in full_text:
            full_text = full_text.replace(search_key, value)
            modified = True

    if modified:
        runs[0].text = full_text
        for run in runs[1:]:
            run.text = ''


def insert_list_into_table(table, list_data):
    if not list_data:
        return
    column_mapping = {}
    for row_idx, row in enumerate(table.rows):
        for col_idx, cell in enumerate(row.cells):
            cell_text = cell.text.strip()
            for key in list_data[0].keys():
                if f'{{{{{key}}}}}' in cell_text:
                    column_mapping[key] = (row_idx, col_idx)
                    cell.text = ''
        if column_mapping:
            break

    if not column_mapping:
        return

    base_row_idx = min(row_idx for row_idx, _ in column_mapping.values())

    for item in list_data:
        if base_row_idx < len(table.rows):
            row = table.rows[base_row_idx]
        else:
            row = table.add_row()

        for key, (r_idx, col_idx) in column_mapping.items():
            if col_idx < len(row.cells):
                row.cells[col_idx].text = item.get(key, '')

        base_row_idx += 1


def run_generated_file(
    request_data: RequestModel,
    doc: DocumentField,
    template_path: str,
    output_path: str,
) -> None:
    word_doc = Document(template_path)
    all_fields = request_data.common_fields.copy()
    all_fields.update(
        {k: v for k, v in doc.fields.items() if not isinstance(v, list)}
    )

    for para in word_doc.paragraphs:
        replace_text_in_runs(para.runs, all_fields)

    for table in word_doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_text_in_runs(para.runs, all_fields)

    for key, value in doc.fields.items():
        if isinstance(value, list):
            for table in word_doc.tables:
                insert_list_into_table(table, value)
    word_doc.save(output_path)


def delete_files(generated_files: list) -> None:
    for file_path in generated_files:
        try:
            os.remove(file_path)
        except OSError:
            pass
